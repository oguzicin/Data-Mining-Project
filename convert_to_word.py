from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
import markdown
import re

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for key, value in kwargs.items():
        tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + key
        element = tcPr.find(tag)
        if element is None:
            element = OxmlElement(key)
            tcPr.append(element)
        
        for k, v in value.items():
            element.set(qn(k), v)

def add_watermark(doc, image_path):
    # Her sayfanın başına logo ekle (header'a eklenir)
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.5))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_logo_to_header(doc, image_path):
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.2))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_word_document(markdown_file, output_file, logo_path=None):
    # Create a new Document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_paragraph_format = title_style.paragraph_format
    title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_paragraph_format = h1_style.paragraph_format
    h1_paragraph_format.space_before = Pt(18)
    h1_paragraph_format.space_after = Pt(12)
    h1_paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_paragraph_format = h2_style.paragraph_format
    h2_paragraph_format.space_before = Pt(12)
    h2_paragraph_format.space_after = Pt(6)
    h2_paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # Normal text style
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_paragraph_format.space_after = Pt(6)
    
    # Read markdown content
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process content
    lines = content.split('\n')
    current_table = None
    in_table = False
    
    for line in lines:
        if line.startswith('# '):  # Title
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):  # H1
            doc.add_heading(line[3:], level=2)
        elif line.startswith('|'):  # Table
            if not in_table:
                # Start new table
                headers = [cell.strip() for cell in line.split('|')[1:-1]]
                current_table = doc.add_table(rows=1, cols=len(headers))
                current_table.style = 'Table Grid'
                header_cells = current_table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                in_table = True
            elif line.startswith('|-'):  # Table separator
                continue
            else:  # Table row
                data = [cell.strip() for cell in line.split('|')[1:-1]]
                row_cells = current_table.add_row().cells
                for i, value in enumerate(data):
                    row_cells[i].text = value
        elif line.strip() == '':
            if in_table:
                in_table = False
                current_table = None
            doc.add_paragraph('', style='CustomNormal')
        else:
            if not in_table:
                doc.add_paragraph(line, style='CustomNormal')
    
    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Add watermark
    if logo_path:
        add_watermark(doc, logo_path)
    
    # Save the document
    doc.save(output_file)

def create_english_report(markdown_file, output_file, logo_path=None):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # Styles
    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(12)
    styles['Heading 1'].font.name = 'Times New Roman'
    styles['Heading 1'].font.size = Pt(14)
    styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.name = 'Times New Roman'
    styles['Heading 2'].font.size = Pt(12)
    styles['Heading 2'].font.bold = True
    # Add logo
    if logo_path:
        add_logo_to_header(doc, logo_path)
    # Read and process markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    in_table = False
    table_data = []
    for line in lines:
        line = line.rstrip('\n')
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('|') and '---' not in line:
            in_table = True
            table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
        elif in_table and (line.strip() == '' or not line.startswith('|')):
            table = doc.add_table(rows=0, cols=len(table_data[0]))
            table.style = 'Table Grid'
            for row in table_data:
                cells = table.add_row().cells
                for i, cell in enumerate(row):
                    cells[i].text = cell
            in_table = False
            table_data = []
            if line.strip() != '':
                doc.add_paragraph(line, style='Normal')
        elif in_table:
            table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
        else:
            doc.add_paragraph(line, style='Normal')
    # Add page number
    add_page_number(doc)
    doc.save(output_file)

if __name__ == '__main__':
    create_word_document('report.docx', 'final_report.docx', logo_path='logo.png')
    create_english_report('report_en.md', 'final_report_en.docx', logo_path='logo.png') 